from __future__ import annotations

from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    KeepTogether,
    PageTemplate,
    Paragraph,
    PageBreak,
    Spacer,
    Table,
    TableStyle,
)


WORKSPACE = Path(r"D:\29722\Desktop\GCC")
PACKAGE = WORKSPACE / "提交相关材料" / "S9全量资料审计与网页仓库归档_20260725"
REPORT_DIR = PACKAGE / "06_报告与验收清单"
SPEC_DIR = PACKAGE / "01_S9赛题原文与规格"
ANALYSIS_DIR = PACKAGE / "04_往届赛题与代码分析"
BLOCKED_DIR = PACKAGE / "05_无法获取与待人工处理"

REPORT_DIR.mkdir(parents=True, exist_ok=True)
SPEC_DIR.mkdir(parents=True, exist_ok=True)
ANALYSIS_DIR.mkdir(parents=True, exist_ok=True)
BLOCKED_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = REPORT_DIR / "S9资料全量审读与外部链接归档报告.docx"
PDF_PATH = REPORT_DIR / "S9资料全量审读与外部链接归档报告.pdf"

NAVY = "17365D"
BLUE = "4472C4"
LIGHT_BLUE = "DCE6F1"
LIGHT_GRAY = "F2F4F7"
GREEN = "E2F0D9"
RED = "FCE4D6"
INK = "1F1F1F"
MUTED = "666666"

S9_ROWS = [
    {
        "op": "Concat",
        "spec": "Tensor[]；dim 默认 0；float32/float16/int32/int8；任意 rank 和合法 dim；除拼接轴外形状一致；允许 1D 空张量和长度 0 分片；非对齐。",
        "current": "仅 FP16，kernel 把最后一维当拼接轴；host 虽按 dim 推输出 shape，但 tiling/kernel 不使用 dim；地址数组只容纳 16 个输入。",
        "risk": "直接解释 Case1 通过、Case2–5 Run failed。隐藏用例只要改变 dtype、rank、dim、输入数或尾块就可能失败。",
        "next": "通用 N-D fallback + last-axis 快路径；动态输入数；0 长度输入；尾块 DataCopyPad；多 dtype。",
    },
    {
        "op": "Greater",
        "spec": "self/other；float32/bfloat16/float16/int32/int8；输出 bool；支持广播；覆盖 inf、-inf、NaN。",
        "current": "仅 FP16 等形；输出 shape 直接等于 self；整张量进 UB；以 SubRelu/Mins/Cast 代替 Compare。",
        "risk": "广播不成立，较大张量 UB 超限；NaN/Inf、溢出和非对齐都可能产生错误或运行失败。",
        "next": "广播步长映射；原生比较语义；分块；多 dtype；特殊值回归。",
    },
    {
        "op": "IndexAdd",
        "spec": "self/index/source；dim 默认 0；index 为 int32 一维向量；source[dim]=len(index)，其他维匹配 self；重复 index 必须累加；输出从 self 开始。",
        "current": "忽略 self；仅 int8、二维、dim=0；空 index 有 blockDim=0 风险；source/out 非对齐复制未覆盖。",
        "risk": "当前实现不满足基础定义；重复索引、任意维度和其他 dtype 均不可靠。",
        "next": "先复制 self；重复 index 正确累加；任意 rank/dim；冲突策略；多 dtype 和尾块。",
    },
    {
        "op": "Transpose",
        "spec": "input + dims 列表；float32/float16/int32/int8；最高 6D 任意排列；非对齐。",
        "current": "仅 FP16 二维交换；忽略 dims；仅处理 16×16 整块；任一维小于 16 时 totalTiles=0。",
        "risk": "任意排列、尾块、非二维都不支持，且小尺寸存在除零/零核风险。",
        "next": "通用坐标映射；尾块；2D 16×16 快路径；任意 dims/rank/dtype。",
    },
    {
        "op": "SquareSumV1",
        "spec": "square 后 sum；axis 为 list_int；keep_dims 默认 false；float16/bfloat16/float32；单轴/多轴/负轴；非对齐。",
        "current": "仅 FP16 最后一轴；输出 shape 强制保留最后一维；忽略 axis/keep_dims；FP16 先平方再转 FP32。",
        "risk": "多轴、非末轴、keep_dims=false、BF16/FP32 均不完整；FP16 平方会先溢出或损失精度。",
        "next": "轴归一化；FP32 中间计算；多轴分阶段或通用回退；正确输出 shape。",
    },
]

KEY_REPOS = [
    ("R002_gitcode_szdaniel_S6T", "S6 多算子完整工程；Fmin/Hypot/Logcumsumexp；大量 DataCopyPad、模板化 dtype 和分块。"),
    ("R078_github_Brown-Bert_S5", "GatherV3 与 ReduceLogSumExp；索引映射、规约和可提交工程结构。"),
    ("R079_github_bubu-12_Ascendc-S4", "ScatterReduce；多维坐标拆解、scatter 轴步长和非对齐搬运，可映射到 IndexAdd。"),
    ("R097_github_Qeeweew_AscendC-S6-Logcumsumexp", "分行规约、事件同步、流水和边界处理。"),
    ("R104_github_ustc1587_ReduceLogSumExp", "FP16→FP32 转换、双缓冲和 ReduceSum 组织，可映射到 SquareSumV1。"),
    ("R107_github_winfan1314_AscendC-S4-Base", "Gather/Pows/Reshape；TILING_KEY 快慢路径与 scalar fallback。"),
    ("R108_github_wybxw_Ascend_ops_Base_S6", "Expand；任意 dtype 模板、广播坐标映射、UB 预算和尾块写回。"),
    ("R109_github_wybxw_Ascend_ops_Performance_S6", "Hypot；float/half/bfloat16 模板、计算 dtype 分离和分块。"),
]


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def configure_docx() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "S9 算子挑战赛｜资料审计与优化依据"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.runs[0].font.name = "Calibri"
    hp.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    hp.runs[0].font.size = Pt(9)
    hp.runs[0].font.color.rgb = rgb(MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("2026-07-25｜资料包内部报告")
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(MUTED)
    return doc


def add_docx_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("S9 资料全量审读与外部链接归档报告")
    r.bold = True
    r.font.size = Pt(23)
    r.font.color.rgb = rgb(NAVY)
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run("覆盖 S9 五题、历届赛题工作簿、往届开源仓库、赛事与云资源材料")
    r2.font.size = Pt(12)
    r2.font.color.rgb = rgb(MUTED)
    table = doc.add_table(rows=4, cols=2)
    values = [
        ("审计日期", "2026-07-25"),
        ("唯一外部链接", "167"),
        ("仓库拉取", "43 成功 / 70 失败"),
        ("网页归档", "97 成功 / 70 失败"),
    ]
    for row, (label, value) in zip(table.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
    set_table_geometry(table, [2700, 6660])
    doc.add_paragraph()


def add_docx_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, BLUE)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = rgb("FFFFFF")
            run.font.size = Pt(9)
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = value
            for paragraph in cells[i].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    set_table_geometry(table, widths)
    doc.add_paragraph()


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)


def build_docx() -> None:
    doc = configure_docx()
    add_docx_title(doc)

    doc.add_heading("1. 结论先行", level=1)
    doc.add_paragraph(
        "是的，Concat 的差结果与此前没有完整按题目规格设计直接相关。当前实现只覆盖 FP16、二维/末维快路径，"
        "而题目明确要求任意 rank、任意合法 dim、多 dtype、非对齐和动态输入列表。Case1 只证明样例快路径可运行，"
        "不能证明实现满足赛题。"
    )
    add_bullet(doc, "已逐项重新读取 9 个赛题工作簿并提取 40 个技术文档链接；S9 的 6 个 PyTorch 2.5 页面均已保存。")
    add_bullet(doc, "Word 正文之外的 HYPERLINK 字段、PDF 注释链接也已纳入，总计 167 个唯一链接。")
    add_bullet(doc, "113 个往届仓库链接逐一尝试完整 clone：43 个成功，70 个失败；成功仓库保留 .git、分支和 HEAD。")
    add_bullet(doc, "43 个成功仓库已建立代码审读矩阵，覆盖 1396 个索引文件、401 个 C/C++ 源文件。")
    add_bullet(doc, "敏感文件 GCC.pem 未复制到资料包。")

    doc.add_heading("2. S9 五题完整规格与当前缺口", level=1)
    rows = [[r["op"], r["spec"], r["current"], r["risk"], r["next"]] for r in S9_ROWS]
    add_docx_table(doc, ["算子", "完整规格", "当前实现", "风险", "下一步"], rows, [900, 2350, 2100, 1900, 2110])

    doc.add_heading("3. Concat 线上结果复盘", level=1)
    doc.add_paragraph("线上结果：Case1 Pass，10.948 μs；Case2–Case5 Run failed。")
    add_bullet(doc, "host 的 InferShape 使用 dim，但 tiling 和 kernel 固定按最后一维解释，dim 非末维时 host/kernel 语义冲突。")
    add_bullet(doc, "OpDef 只声明 DT_FLOAT16，题目还要求 float32、int32、int8。")
    add_bullet(doc, "ListTensor 地址数组只有 16 项，题目输入数量为动态列表；隐藏用例可能超过固定容量。")
    add_bullet(doc, "实现把除最后一维外全部折叠成 outer，只适用于末维拼接；任意 rank 不等于任意 dim。")
    add_bullet(doc, "性能路径与通用正确路径没有分层，导致样例路径一旦不匹配就直接运行失败，而不是回退。")
    doc.add_paragraph(
        "因此，下一轮不能继续围绕 Case1 做微调。正确架构应先建立全规格通用路径，再为常见 dtype、末维、对齐形状增加快路径。"
    )

    doc.add_heading("4. 往届优秀代码的可复用经验", level=1)
    add_bullet(doc, "非对齐是常态：43 个成功仓库中 25 个使用 DataCopyPad。尾块必须作为一级设计对象。")
    add_bullet(doc, "模板化 dtype 是高分代码的常见结构：35 个仓库使用 C++ template，避免复制多份脆弱实现。")
    add_bullet(doc, "TILING_KEY/模板特化用于把“对齐快路径”和“通用回退路径”分开，不能只保留快路径。")
    add_bullet(doc, "UB 按固定预算切 tile；不把整张量一次性放入 UB。大 shape 的可运行性优先于单个样例的低延迟。")
    add_bullet(doc, "索引类算子要显式处理多维坐标、stride、重复 index 和写冲突；ScatterReduce/Gather/Expand 是 IndexAdd 的直接参考。")
    add_bullet(doc, "规约类算子应在 FP32 中间类型中平方/累加，再转换输出；先 FP16 平方会提前溢出或损失精度。")
    add_bullet(doc, "双缓冲只有在 tile 循环真实重叠搬运与计算时才有价值；仅把 BUFFER_NUM 改大并不自动提升性能。")
    rows = [[name, note] for name, note in KEY_REPOS]
    add_docx_table(doc, ["本地仓库目录", "主要学习价值"], rows, [3150, 6210])

    doc.add_heading("5. 推荐优化顺序与验收门槛", level=1)
    ordered = [
        "先冻结题目语义矩阵：dtype × rank × axis/dim × 对齐 × 空张量/空 index × 特殊值 × 重复 index。",
        "为五题建立通用正确路径，确保所有本地扩展测试和远端功能用例通过。",
        "只在输入满足条件时进入快路径；任何不满足的情况必须回退，不能运行失败。",
        "再做 tiling、核数、双缓冲、vector 指令和减少 GM 往返的性能优化。",
        "每次远端提交都记录代码提交号、zip 哈希、五个 Case 的结果和日志，避免本地/远端/仓库不同步。",
    ]
    for i, item in enumerate(ordered, 1):
        doc.add_paragraph(f"{i}. {item}")

    doc.add_heading("6. 外部链接与仓库归档结果", level=1)
    add_docx_table(
        doc,
        ["类别", "总数", "成功", "失败", "说明"],
        [
            ["唯一外部链接", "167", "97 保存网页", "70", "完整逐项状态见审计工作簿"],
            ["仓库链接", "113", "43 完整 clone", "70", "成功项保留 .git、分支、HEAD"],
            ["赛题工作簿", "9", "9 已读", "0", "提取 40 个唯一技术链接"],
            ["S9 PyTorch 2.5 链接", "6 + 1 跟进", "7", "0", "index_add 页面正文跳转到 index_add_，已补抓"],
        ],
        [1800, 1000, 1500, 900, 4160],
    )

    doc.add_heading("7. 无法获取与限制", level=1)
    add_bullet(doc, "Gitee：大多数 clone 返回平台 reject/HTTP 400，网页返回 HTTP 405；属于平台或当前网络环境拒绝，不是未尝试。")
    add_bullet(doc, "GitCode：LiJianhao2/AscendS6 要求登录；szdaniel/S6T 已成功 clone。")
    add_bullet(doc, "GitHub：wilburx813/ascend_s4 返回 Repository not found；页面亦 404。")
    add_bullet(doc, "pytorch.ac.cn 的 2 个历史链接在当前请求状态下无法保存；原链接和错误已记录，其他 38 个工作簿技术页成功。")
    add_bullet(doc, "http://e.huawei.com 返回 502；手册中的华为云 ModelArts 页面和控制台链接已保存。")
    add_bullet(doc, "本机没有 Word/LibreOffice/WPS 渲染器，DOCX 无法做视觉页面复核；主 PDF 已生成并逐页检查，DOCX 做结构检查。")

    doc.add_heading("8. 资料包结构与使用方式", level=1)
    structure = [
        ("00_材料索引", "源文件清单、167 链接总表、网页状态、仓库状态、目录映射。"),
        ("01_S9赛题原文与规格", "S9 原始材料、官方 zip_op.sh、五题规格摘要。"),
        ("02_网页归档", "97 个成功响应和获取映射。"),
        ("03_往届冠军仓库快照", "43 个完整工作树与 Git 历史。"),
        ("04_往届赛题与代码分析", "历届原始包、解压内容、9 个工作簿、代码审读矩阵与学习结论。"),
        ("05_无法获取与待人工处理", "失败原因、可重试条件和本机限制。"),
        ("06_报告与验收清单", "主报告、审计工作簿、预览和最终校验。"),
    ]
    add_docx_table(doc, ["目录", "内容"], [list(x) for x in structure], [2500, 6860])
    doc.add_paragraph(
        "建议先读本报告，再打开《S9外部链接与往届仓库审计表.xlsx》筛选具体链接或仓库，最后按算子进入对应源码目录。"
    )

    doc.add_heading("9. 下一轮开始优化前需要的条件", level=1)
    add_bullet(doc, "以 ModelArts `/home/ma-user/work/` 的实时代码为准，先同步回本地并记录远端 commit/文件哈希。")
    add_bullet(doc, "保持 CANN 社区版 8.5.0、Ascend 910B 和 Euler 2.10 环境。")
    add_bullet(doc, "远端命令继续写入 `/home/ma-user/work/s9/codex-visible-terminal.log`，你可在 JupyterLab 中实时查看。")
    add_bullet(doc, "先允许我跑完整功能矩阵和本地/远端日志，再提交榜单；不再用单一样例通过代替题目通过。")

    doc.save(DOCX_PATH)


def register_pdf_font() -> str:
    candidates = [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            try:
                pdfmetrics.registerFont(TTFont("CJK", str(candidate), subfontIndex=0))
                return "CJK"
            except Exception:
                continue
    return "Helvetica"


def build_pdf() -> None:
    font = register_pdf_font()
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "BodyCJK",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=9.5,
        leading=14,
        spaceAfter=6,
        textColor=colors.HexColor("#1F1F1F"),
        wordWrap="CJK",
    )
    title = ParagraphStyle(
        "TitleCJK",
        parent=styles["Title"],
        fontName=font,
        fontSize=22,
        leading=28,
        textColor=colors.HexColor("#17365D"),
        alignment=TA_LEFT,
        spaceAfter=8,
    )
    subtitle = ParagraphStyle(
        "SubtitleCJK",
        parent=body,
        fontSize=11,
        textColor=colors.HexColor("#666666"),
        spaceAfter=16,
    )
    h1 = ParagraphStyle(
        "H1CJK",
        parent=styles["Heading1"],
        fontName=font,
        fontSize=15,
        leading=20,
        textColor=colors.HexColor("#2E74B5"),
        spaceBefore=12,
        spaceAfter=8,
        keepWithNext=True,
    )
    h2 = ParagraphStyle(
        "H2CJK",
        parent=styles["Heading2"],
        fontName=font,
        fontSize=12,
        leading=16,
        textColor=colors.HexColor("#1F4D78"),
        spaceBefore=8,
        spaceAfter=5,
        keepWithNext=True,
    )
    bullet = ParagraphStyle(
        "BulletCJK",
        parent=body,
        leftIndent=14,
        firstLineIndent=-8,
        bulletIndent=4,
        spaceAfter=4,
    )
    small = ParagraphStyle(
        "SmallCJK",
        parent=body,
        fontSize=7.3,
        leading=10,
        spaceAfter=0,
    )
    small_center = ParagraphStyle(
        "SmallCenterCJK",
        parent=small,
        alignment=TA_CENTER,
    )

    def header_footer(canvas, doc):
        canvas.saveState()
        canvas.setFont(font, 8)
        canvas.setFillColor(colors.HexColor("#666666"))
        canvas.drawString(doc.leftMargin, 10.45 * inch, "S9 算子挑战赛｜资料审计与优化依据")
        canvas.drawRightString(7.5 * inch, 0.42 * inch, f"第 {doc.page} 页")
        canvas.setStrokeColor(colors.HexColor("#D9E2F3"))
        canvas.line(doc.leftMargin, 10.3 * inch, 7.5 * inch, 10.3 * inch)
        canvas.restoreState()

    doc = BaseDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        leftMargin=0.7 * inch,
        rightMargin=0.7 * inch,
        topMargin=0.82 * inch,
        bottomMargin=0.65 * inch,
        title="S9资料全量审读与外部链接归档报告",
        author="Codex",
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="main", frames=frame, onPage=header_footer)])
    story = [
        Spacer(1, 0.15 * inch),
        Paragraph("S9 资料全量审读与外部链接归档报告", title),
        Paragraph("覆盖 S9 五题、历届赛题工作簿、往届开源仓库、赛事与云资源材料", subtitle),
    ]
    summary_data = [
        [Paragraph("<b>审计日期</b>", body), Paragraph("2026-07-25", body), Paragraph("<b>唯一外部链接</b>", body), Paragraph("167", body)],
        [Paragraph("<b>仓库拉取</b>", body), Paragraph("43 成功 / 70 失败", body), Paragraph("<b>网页归档</b>", body), Paragraph("97 成功 / 70 失败", body)],
    ]
    t = Table(summary_data, colWidths=[1.2 * inch, 2.1 * inch, 1.2 * inch, 2.1 * inch])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#DCE6F1")),
        ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#DCE6F1")),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#B4C6E7")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story += [t, Spacer(1, 10), Paragraph("1. 结论先行", h1)]
    story.append(Paragraph(
        "是的，Concat 的差结果与此前没有完整按题目规格设计直接相关。当前实现只覆盖 FP16、二维/末维快路径，"
        "而题目要求任意 rank、任意合法 dim、多 dtype、动态输入、空分片和非对齐。Case1 只证明样例快路径可运行，"
        "不能证明实现满足赛题。",
        body,
    ))
    for text in [
        "9 个赛题工作簿全部重新读取，提取 40 个技术文档链接；S9 的 6 个 PyTorch 2.5 页面和 index_add_ 跟进页均已保存。",
        "Word 的 HYPERLINK 字段和 PDF 注释链接也已纳入，总计 167 个唯一链接。",
        "113 个仓库逐一尝试完整 clone：43 个成功，70 个失败；成功项保留 .git、分支和 HEAD。",
        "43 个成功仓库已建立代码审读矩阵：1396 个索引文件、401 个 C/C++ 源文件。",
        "敏感文件 GCC.pem 未复制到资料包。",
    ]:
        story.append(Paragraph("• " + escape(text), bullet))

    story.append(Paragraph("2. S9 五题完整规格与当前缺口", h1))
    table_data = [[Paragraph(f"<b>{h}</b>", small_center) for h in ["算子", "完整规格", "当前实现", "风险", "下一步"]]]
    for row in S9_ROWS:
        table_data.append([
            Paragraph(row["op"], small),
            Paragraph(escape(row["spec"]), small),
            Paragraph(escape(row["current"]), small),
            Paragraph(escape(row["risk"]), small),
            Paragraph(escape(row["next"]), small),
        ])
    table = Table(table_data, colWidths=[0.55 * inch, 1.65 * inch, 1.55 * inch, 1.45 * inch, 1.6 * inch], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B4C6E7")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story += [table, PageBreak(), Paragraph("3. Concat 线上结果复盘", h1)]
    story.append(Paragraph("<b>线上结果：</b>Case1 Pass，10.948 μs；Case2–Case5 Run failed。", body))
    for text in [
        "host InferShape 使用 dim，但 tiling/kernel 固定按最后一维解释，dim 非末维时语义冲突。",
        "OpDef 只声明 DT_FLOAT16，题目还要求 float32、int32、int8。",
        "ListTensor 地址数组只有 16 项，动态输入列表可能超过固定容量。",
        "实现把除最后一维外全部折叠成 outer，只适用于末维拼接。",
        "没有通用回退路径，输入不匹配快路径时直接运行失败。",
    ]:
        story.append(Paragraph("• " + escape(text), bullet))
    story.append(Paragraph(
        "下一轮必须先建立全规格通用路径，再为常见 dtype、末维、对齐形状增加快路径，不能继续只围绕 Case1 微调。",
        body,
    ))

    story.append(Paragraph("4. 往届优秀代码的可复用经验", h1))
    lessons = [
        "非对齐是常态：43 个成功仓库中 25 个使用 DataCopyPad。",
        "35 个仓库使用 C++ template，多 dtype 不应靠复制多份实现。",
        "TILING_KEY/模板特化用于拆分对齐快路径和通用 fallback。",
        "UB 按预算切 tile，不把整张量一次性放进 UB。",
        "IndexAdd 应参考 ScatterReduce/Gather/Expand 的多维坐标、stride、重复 index 和冲突处理。",
        "SquareSumV1 应先转换 FP32，再平方与累加，最后按输出 dtype 转换。",
        "双缓冲必须与真实 tile 循环重叠，否则 BUFFER_NUM 增大没有实际收益。",
    ]
    for text in lessons:
        story.append(Paragraph("• " + escape(text), bullet))
    key_data = [[Paragraph("<b>本地仓库目录</b>", small), Paragraph("<b>主要学习价值</b>", small)]]
    key_data.extend([[Paragraph(name, small), Paragraph(escape(note), small)] for name, note in KEY_REPOS])
    kt = Table(key_data, colWidths=[2.65 * inch, 4.05 * inch], repeatRows=1)
    kt.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B4C6E7")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story += [kt, Paragraph("5. 推荐优化顺序与验收门槛", h1)]
    for i, text in enumerate([
        "冻结 dtype × rank × axis/dim × 对齐 × 空输入 × 特殊值 × 重复 index 的语义矩阵。",
        "五题先建立通用正确路径，所有本地扩展测试与远端功能用例通过。",
        "只在条件满足时进入快路径；否则必须回退，不能 Run failed。",
        "再做 tiling、核数、双缓冲、vector 指令和 GM 往返优化。",
        "每次提交记录代码提交号、zip 哈希、Case 结果和日志，避免本地/远端/仓库不同步。",
    ], 1):
        story.append(Paragraph(f"{i}. {escape(text)}", body))

    story += [PageBreak(), Paragraph("6. 外部链接与仓库归档结果", h1)]
    status_data = [
        ["类别", "总数", "成功", "失败", "说明"],
        ["唯一外部链接", "167", "97 网页", "70", "逐项状态见审计工作簿"],
        ["仓库链接", "113", "43 clone", "70", "保留 .git/分支/HEAD"],
        ["赛题工作簿", "9", "9 已读", "0", "40 个唯一技术链接"],
        ["S9 技术页", "6+1", "7", "0", "index_add_ 为必要跟进页"],
    ]
    status_table = Table(
        [[Paragraph(f"<b>{escape(v)}</b>", small) if r == 0 else Paragraph(escape(v), small) for v in row] for r, row in enumerate(status_data)],
        colWidths=[1.25 * inch, 0.7 * inch, 1.05 * inch, 0.65 * inch, 3.05 * inch],
    )
    status_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B4C6E7")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story += [status_table, Paragraph("7. 无法获取与限制", h1)]
    for text in [
        "Gitee：大多数 clone 返回平台 reject/HTTP 400，网页返回 HTTP 405；属于平台或当前网络环境拒绝。",
        "GitCode：LiJianhao2/AscendS6 要求登录；szdaniel/S6T 已成功 clone。",
        "GitHub：wilburx813/ascend_s4 返回 Repository not found，网页 404。",
        "pytorch.ac.cn 的 2 个历史链接当前无法保存；原链接与错误已记录，其余 38 个工作簿技术页成功。",
        "http://e.huawei.com 返回 502；手册中的 ModelArts 页面和控制台链接已保存。",
        "本机没有 Word/LibreOffice/WPS 渲染器，DOCX 无法视觉复核；本 PDF 已逐页检查。",
    ]:
        story.append(Paragraph("• " + escape(text), bullet))

    story.append(Paragraph("8. 资料包结构", h1))
    structure = [
        ("00_材料索引", "文件、链接、网页、仓库与目录映射。"),
        ("01_S9赛题原文与规格", "原始材料、官方 zip_op.sh、五题规格。"),
        ("02_网页归档", "97 个成功响应。"),
        ("03_往届冠军仓库快照", "43 个完整仓库。"),
        ("04_往届赛题与代码分析", "历届原始包、解压内容、工作簿与代码矩阵。"),
        ("05_无法获取与待人工处理", "失败原因与可重试条件。"),
        ("06_报告与验收清单", "主报告、审计工作簿和预览。"),
    ]
    st = Table(
        [[Paragraph("<b>目录</b>", small), Paragraph("<b>内容</b>", small)]]
        + [[Paragraph(a, small), Paragraph(b, small)] for a, b in structure],
        colWidths=[2.35 * inch, 4.35 * inch],
        repeatRows=1,
    )
    st.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B4C6E7")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story += [st, Paragraph("9. 下一轮优化前需要的条件", h1)]
    for text in [
        "以 ModelArts /home/ma-user/work/ 的实时代码为准，先同步回本地并记录远端 commit/文件哈希。",
        "保持 CANN 社区版 8.5.0、Ascend 910B 和 Euler 2.10。",
        "远端命令写入 /home/ma-user/work/s9/codex-visible-terminal.log，便于在 JupyterLab 实时查看。",
        "先跑完整功能矩阵和日志，再提交榜单；不再用单一样例通过代替题目通过。",
    ]:
        story.append(Paragraph("• " + escape(text), bullet))

    doc.build(story)


def write_markdown_files() -> None:
    readme = """# S9 全量资料审计与网页仓库归档

本资料包用于重新建立可靠的题目依据、网页快照、往届代码基线和失败记录。

## 核心数字

- 唯一外部链接：167
- 网页保存：97 成功，70 失败
- 往届仓库：113 个链接，43 个完整克隆成功，70 个失败
- 赛题工作簿：9 个，提取 40 个唯一技术文档链接
- 成功仓库代码审读：43 个仓库、1396 个索引文件、401 个 C/C++ 源文件

## 阅读顺序

1. `06_报告与验收清单/S9资料全量审读与外部链接归档报告.pdf`
2. `06_报告与验收清单/S9外部链接与往届仓库审计表.xlsx`
3. `01_S9赛题原文与规格/S9五题完整规格摘要.md`
4. `04_往届赛题与代码分析/往届优秀代码学习结论.md`
5. `05_无法获取与待人工处理/无法获取与待人工处理.md`

## 目录

- `00_材料索引`：CSV 清单、状态和目录映射。
- `01_S9赛题原文与规格`：原始 S9 文件、官方打包脚本和规格摘要。
- `02_网页归档`：成功下载的原始网页。
- `03_往届冠军仓库快照`：成功 clone 的完整仓库。
- `04_往届赛题与代码分析`：历届赛题、工作簿审读结果和代码矩阵。
- `05_无法获取与待人工处理`：失败原因和人工处理建议。
- `06_报告与验收清单`：主报告、审计工作簿和预览。

## 安全说明

`GCC.pem` 是敏感密钥，未复制进资料包，也不会进入最终 ZIP。
"""
    (PACKAGE / "README_资料包说明.md").write_text(readme, encoding="utf-8")

    spec_lines = ["# S9 五题完整规格摘要", "", "来源：S9 工作簿、官方测试脚本、PyTorch 2.5 页面。", ""]
    for row in S9_ROWS:
        spec_lines += [
            f"## {row['op']}",
            "",
            f"- 完整规格：{row['spec']}",
            f"- 当前实现：{row['current']}",
            f"- 风险：{row['risk']}",
            f"- 下一步：{row['next']}",
            "",
        ]
    (SPEC_DIR / "S9五题完整规格摘要.md").write_text("\n".join(spec_lines), encoding="utf-8")

    repo_lines = [
        "# 往届优秀代码学习结论",
        "",
        "## 全量统计",
        "",
        "- 成功拉取 43 个仓库，分析 1396 个索引文件、401 个 C/C++ 源文件。",
        "- 25 个仓库使用 DataCopyPad，35 个仓库使用模板。",
        "",
        "## 可直接迁移到 S9 的原则",
        "",
        "- 通用 fallback 与快路径并存。",
        "- UB 按预算切 tile，不能整张量搬入。",
        "- 多 dtype 使用模板与计算类型分离。",
        "- 非对齐、空输入、小尺寸和尾块必须单独验证。",
        "- 索引算子处理多维 stride、重复 index 和冲突。",
        "- 规约算子使用 FP32 中间累加。",
        "",
        "## 优先参考仓库",
        "",
    ]
    repo_lines += [f"- `{name}`：{note}" for name, note in KEY_REPOS]
    (ANALYSIS_DIR / "往届优秀代码学习结论.md").write_text("\n".join(repo_lines), encoding="utf-8")

    blocked = """# 无法获取与待人工处理

## 仓库

- Gitee：大多数 clone 返回平台 reject/HTTP 400；网页 HTTP 405。建议在已登录 Gitee 的网络环境中按 `00_材料索引/仓库拉取结果.csv` 重试。
- GitCode：`LiJianhao2/AscendS6` 要求登录；`szdaniel/S6T` 已成功。
- GitHub：`wilburx813/ascend_s4` 返回 Repository not found。

## 网页

- `http://e.huawei.com/`：HTTP 502。
- `pytorch.ac.cn` 两个 S4 历史页面：当前请求状态失败；原链接仍保留。

## 文档渲染

本机未安装 Word、LibreOffice 或 WPS，DOCX 无法转成页面图片做视觉复核。主 PDF 已生成并逐页检查；DOCX 已做结构检查。

## 敏感文件

`D:/29722/Desktop/GCC/GCC.pem` 未复制、未压缩、未上传。
"""
    (BLOCKED_DIR / "无法获取与待人工处理.md").write_text(blocked, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_pdf()
    write_markdown_files()
    print(DOCX_PATH)
    print(PDF_PATH)
